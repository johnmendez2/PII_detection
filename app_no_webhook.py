import json
import time
import threading
import requests
from uuid import uuid4
import uuid
import datetime
from dotenv import load_dotenv
from utils.version import API_VERSION, SERVICE_NAME
from utils.status_codes import StatusCodes
from api.models import response_template
import re
from flask_swagger_ui import get_swaggerui_blueprint
from utils.s3_download_file import download_s3_files
from utils.s3_presign import create_presigned_url
from utils.s3_upload_object import upload_to_s3
from utils.s3_presigned_download import generate_presigned_url
from utils.s3_delete_object import generate_presigned_delete, generate_presigned_delete_for_folder
from flask import Flask, request, jsonify, send_file, make_response
from PyPDF2 import PdfReader
from docx import Document
import tempfile
import os
from werkzeug.utils import secure_filename
import json
import argparse
from itertools import chain
import pandas as pd
from pathlib import Path
from transformers import (
    AutoTokenizer,
    AutoModelForTokenClassification,
    Trainer,
    TrainingArguments,
    DataCollatorForTokenClassification,
)
from datasets import Dataset
import numpy as np
from spire.pdf.common import *
from spire.pdf import *
from spire.doc import *
from spire.doc.common import *
from docx import Document
from spacy.lang.en import English
from PyPDF2 import PdfReader, PdfWriter
import zipfile
import shutil

app = Flask(__name__)
# Load environment variables
load_dotenv(override=True)
app.config.from_object(__name__)  # Load config from object

# Load JSON configuration
with open("config.json") as f:
    config = json.load(f)

# Now you can access your configuration values like this:
# debug_mode = app.config["DEBUG"]

############### SWAGGER DOCUMENTATION ###############

SWAGGER_URL = "/api/docs"  # URL for exposing Swagger UI (without trailing '/')
cwd = os.getcwd()
API_URL = "/static/swagger.json"  # Our API url (can of course be a local resource)
# Call factory function to create our blueprint
swaggerui_blueprint = get_swaggerui_blueprint(
    SWAGGER_URL,  # Swagger UI static files will be mapped to '{SWAGGER_URL}/dist/'
    API_URL,
    config={"app_name": "PII DETECTION"},  # Swagger UI config overrides
)

app.register_blueprint(swaggerui_blueprint)

############### ENV VARIABLES ###############
SUPPORTED_METHOD = ["upload_files", "process_and_extract"]


############### ADD YOUR AI MARKETPLACE WEBHOOK ENDPOINT HERE ###############
webhook_url = "http://localhost:8000/callback"


############### ADD YOUR CUSTOM AI AGENT CALL HERE ###############
def process_and_extract(user_id):
    tmp_folder_path = os.path.join(os.getcwd(), str(user_id))
    download_s3_files(user_id)
    data = {}
    
    # Iterate over all files in the directory
    for filename in os.listdir(tmp_folder_path):
        file_path = os.path.join(tmp_folder_path, filename)
        
        # Check if it's a file (not a directory)
        if os.path.isfile(file_path):
            # Determine the file type
            if file_path.endswith('.docx'):
                result = extract_text_from_docx(file_path)  # Assume this function exists
                data[secure_filename(filename)] = result
            elif file_path.endswith('.pdf'):
                result = extract_text_from_pdf(file_path)  # Assume this function exists
                data[secure_filename(filename)] = result
            else:
                print(f"Unsupported file type: {filename}")
                continue
    print(data)            
    def tokenize(example, tokenizer):
        text = []
        token_map = []

        idx = 0

        for t, ws in zip(example["tokens"], example["trailing_whitespace"]):
            text.append(t)
            token_map.extend([idx] * len(t))
            if ws:
                text.append(" ")
                token_map.append(-1)

            idx += 1

        tokenized = tokenizer(
            "".join(text),
            return_offsets_mapping=True,
            truncation=False,
            max_length=INFERENCE_MAX_LENGTH,
            stride=STRIDE,
            return_overflowing_tokens=True,
        )

        return {
            **tokenized,
            "token_map": token_map,
        }

    # Define your tokenizer (replace 'bert-base-uncased' with your desired tokenizer)
    tokenizer = English().tokenizer
    dataset = []

    # Iterate over the values of data, which are dictionaries
    for item in data.values():
        tokens = tokenizer(item["text"])
        dataset.append(
            {
                "document": item["filename"],
                "full_text": item["text"],
                "tokens": [token.text for token in tokens],
                "trailing_whitespace": [
                    True if token.whitespace_ == " " else False for token in tokens
                ],
                "labels": ["O" for token in tokens],
            }
        )

    def merge_rows(df):
        new_data = []
        merge = False
        for _, row in df.iterrows():
            if row["tokens"] == "<":
                merge = True
                temp = [
                    row["document"],
                    row["tokens"],
                    row["labels"],
                    row["trailing_whitespace"],
                ]
            elif row["tokens"] == ">":
                merge = False
                temp[1] += row["tokens"]
                temp[2] = row["labels"]
                temp[3] = row["trailing_whitespace"]
                new_data.append(temp)
            elif merge:
                temp[1] += row["tokens"]
            else:
                new_data.append(row.tolist())
        return pd.DataFrame(new_data, columns=df.columns)

    datafr = pd.DataFrame(
        {
            "document": [
                doc
                for doc_lst in [
                    [dataset[i]["document"] for _ in range(len(dataset[i]["tokens"]))]
                    for i in range(len(dataset))
                ]
                for doc in doc_lst
            ],
            "tokens": [token for x in dataset for token in x["tokens"]],
            "labels": [token for x in dataset for token in x["labels"]],
            "trailing_whitespace": [
                token for x in dataset for token in x["trailing_whitespace"]
            ],
        }
    )

    datafr = merge_rows(datafr)

    # Assuming dataset is already populated as shown in your previous example

    # Initialize an empty dictionary to hold the aggregated data
    data_dict = {
        "full_text": [],
        "document": [],
        "tokens": [],
        "trailing_whitespace": [],
    }

    # Iterate over the dataset list
    for item in dataset:
        # Append the values to the corresponding lists in data_dict
        data_dict["full_text"].append(item["full_text"])
        data_dict["document"].append(item["document"])
        data_dict["tokens"].append(item["tokens"])
        data_dict["trailing_whitespace"].append(item["trailing_whitespace"])

    # Create a dataset
    ds = Dataset.from_dict(data_dict)
    ds[0]
    # # Display dataset info
    tokenizer = AutoTokenizer.from_pretrained(model_path)
    ds = ds.map(tokenize, fn_kwargs={"tokenizer": tokenizer}, num_proc=2)

    model = AutoModelForTokenClassification.from_pretrained(model_path)
    collator = DataCollatorForTokenClassification(tokenizer, pad_to_multiple_of=16)
    args = TrainingArguments(
        ".",
        per_device_eval_batch_size=1,
        report_to="none",
    )
    trainer = Trainer(
        model=model,
        args=args,
        data_collator=collator,
        tokenizer=tokenizer,
    )

    def backwards_map_preds(sub_predictions, max_len):
        if (
            max_len != 1
        ):  # nothing to map backwards if sequence is too short to be split in the first place
            if i == 0:
                # First sequence needs no SEP token (used to end a sequence)
                sub_predictions = sub_predictions[:, :-1, :]
            elif i == max_len - 1:
                # End sequence needs to CLS token + Stride tokens
                sub_predictions = sub_predictions[
                    :, 1 + STRIDE :, :
                ]  # CLS tokens + Stride tokens
            else:
                # Middle sequence needs to CLS token + Stride tokens + SEP token
                sub_predictions = sub_predictions[:, 1 + STRIDE : -1, :]
        return sub_predictions

    def backwards_map_(row_attribute, max_len):
        # Same logics as for backwards_map_preds - except lists instead of 3darray
        if max_len != 1:
            if i == 0:
                row_attribute = row_attribute[:-1]
            elif i == max_len - 1:
                row_attribute = row_attribute[1 + STRIDE :]
            else:
                row_attribute = row_attribute[1 + STRIDE : -1]
        return row_attribute

    preds = []
    ds_dict = {"document": [], "token_map": [], "offset_mapping": [], "tokens": []}

    for row in ds:
        # keys that need to be re-assembled
        row_preds = []
        row_offset = []

        for i, y in enumerate(row["offset_mapping"]):
            # create new datasset for each of of the splits per document
            x = Dataset.from_dict(
                {
                    "token_type_ids": [row["token_type_ids"][i]],
                    "input_ids": [row["input_ids"][i]],
                    "attention_mask": [row["attention_mask"][i]],
                    "offset_mapping": [row["offset_mapping"][i]],
                }
            )
            # predict for that split
            pred = trainer.predict(x).predictions
            # removing the stride and additional CLS & SEP that are created
            row_preds.append(backwards_map_preds(pred, len(row["offset_mapping"])))
            row_offset += backwards_map_(y, len(row["offset_mapping"]))

        # Finalize row
        ds_dict["document"].append(row["document"])
        ds_dict["tokens"].append(row["tokens"])
        ds_dict["token_map"].append(row["token_map"])
        ds_dict["offset_mapping"].append(row_offset)

        # Finalize prediction collection by concattenating
        p_concat = np.concatenate(row_preds, axis=1)
        preds.append(p_concat)

    config = json.load(open(Path(model_path) / "config.json"))
    id2label = config["id2label"]

    preds_final = []
    for predictions in preds:
        predictions_softmax = np.exp(predictions) / np.sum(
            np.exp(predictions), axis=2
        ).reshape(predictions.shape[0], predictions.shape[1], 1)
        predictions = predictions.argmax(-1)
        predictions_without_O = predictions_softmax[
            :, :, : config["label2id"]["O"]
        ].argmax(-1)
        O_predictions = predictions_softmax[:, :, config["label2id"]["O"]]

        threshold = 0.9
        preds_final.append(
            np.where(O_predictions < threshold, predictions_without_O, predictions)
        )

    ds = Dataset.from_dict(ds_dict)
    pairs = []
    document, token, label, token_str = [], [], [], []
    for p, token_map, offsets, tokens, doc in zip(
        preds_final, ds["token_map"], ds["offset_mapping"], ds["tokens"], ds["document"]
    ):
        for token_pred, (start_idx, end_idx) in zip(p[0], offsets):
            label_pred = id2label[str(token_pred)]

            if start_idx + end_idx == 0:
                continue

            if token_map[start_idx] == -1:
                start_idx += 1

            if start_idx >= len(token_map):
                break

            token_id = token_map[start_idx]

            pair = (doc, token_id)

            if pair not in pairs:
                document.append(doc)
                token.append(token_id)
                label.append(label_pred)
                token_str.append(tokens[token_id])
                pairs.append(pair)

    df = pd.DataFrame(
        {"document": document, "token": token, "label": label, "token_str": token_str}
    )
    df["row_id"] = list(range(len(df)))
    # df

    # Assuming df and datafr are your DataFrames

    # Merge the label information from df into datafr based on document and token_str
    def merge_labels(row):
        matching_rows = df[
            (df["document"] == row["document"]) & (df["token_str"] == row["tokens"])
        ]
        if not matching_rows.empty:
            return matching_rows["label"].values[0]
        elif not row["tokens"].startswith("\\"):
            return "O"
        else:
            return None

    datafr["labels"] = datafr.apply(merge_labels, axis=1)
    print(datafr)
    presign_url, delete_files_url = replace_tokens_in_files(datafr, user_id)
    print(presign_url, delete_files_url)
    return presign_url, delete_files_url


############### MAIN FUNCTIONS ###############
INFERENCE_MAX_LENGTH = 1024
STRIDE = 384
model_path = os.path.join(os.getcwd(), "trainer1400")


def replace_tokens_in_files(data, user_id):
    # Define the path to the tmp folder
    tmp_folder_path = os.path.join(os.getcwd(), user_id)

    # Ensure the tmp folder exists
    if not os.path.exists(tmp_folder_path):
        os.makedirs(tmp_folder_path, exist_ok=True)

    unique_file_paths = data["document"].unique()

    for file_path in unique_file_paths:
        dataframe = data[(data["labels"] != "O") & (data["document"] == file_path)]
        # Check if the file is a DOCX or a PDF
        if file_path.endswith(".docx"):
            doc = Document(file_path)
            unique_tokens = dataframe["tokens"].unique()
            for para in doc.paragraphs:
                for token in unique_tokens:
                    try:
                        para.text = para.text.replace(token, "xxx")
                    except Exception as e:
                        print(
                            f"Error replacing token '{token}': {e}. Skipping this token."
                        )
            # Save the modified DOCX file to the tmp folder
            doc.save(
                os.path.join(
                    tmp_folder_path, f"{file_path.replace('.docx', '')}removedPII.docx"
                )
            )
        elif file_path.endswith(".pdf"):
            doc = PdfDocument()
            doc.LoadFromFile(file_path)
            unique_tokens = dataframe["tokens"].unique()
            for i in range(doc.Pages.Count):
                page = doc.Pages[i]
                replacer = PdfTextReplacer(page)
                for token in unique_tokens:
                    try:
                        replacer.ReplaceAllText(token, "xxx")
                    except Exception as e:
                        print(
                            f"Error replacing token '{token}': {e}. Skipping this token."
                        )
            # Save the modified PDF file to the tmp folder
            doc.SaveToFile(
                os.path.join(
                    tmp_folder_path, f"{file_path.replace('.pdf', '')}removedPII.pdf"
                )
            )
            doc.Close()
        else:
            print(
                f"Unsupported file type for {file_path}. Only PDF and DOCX files are supported."
            )

    # Zip only the contents of the tmp folder that contain 'removedPII' in their name
    zip_path = os.path.join(tmp_folder_path, "modified_files.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(tmp_folder_path):
            for file in files:
                if (
                    "removedPII" in file
                ):  # Include files that contain 'removedPII' anywhere in their name
                    zipf.write(
                        os.path.join(root, file),
                        os.path.relpath(os.path.join(root, file), tmp_folder_path),
                    )
    # Delete all files in tmp_folder_path except for the ZIP file
    for root, dirs, files in os.walk(tmp_folder_path):
        for file in files:
            if file!= "modified_files.zip":
                os.remove(os.path.join(root, file))

    if upload_to_s3(user_id, zip_path):
        print("ZIP file uploaded successfully.")
    else:
        print("Failed to upload ZIP file.")
        return "Failed to upload ZIP file."

    # Delete the tmp_folder_path
    shutil.rmtree(tmp_folder_path)
    print("Temporary folder deleted.")
    presign_download_url = generate_presigned_url(user_id)
    delete_files_url = generate_presigned_delete_for_folder(user_id + "/")
    return presign_download_url, delete_files_url



############### MIDDLEWARE FUNCTIONS ###############
# Modify the return statement in process_query to handle generator functions correctly
def process_query(user_id):
    start_time = time.time()
    download_url, delete_files_url = process_and_extract(user_id)
    # print(f"response_data = {collected_data}")
    # Now, collected_data contains all the data yielded by the generate function
    # You can proceed with your logic to return these data as needed
    # For example, return them as part of a JSON response
    end_time = time.time()
    processing_duration = (
        end_time - start_time
    )  # Calculate processing duration in seconds
    return download_url,delete_files_url, processing_duration


def success_response(task_id, data, requestId, trace_id, process_duration):
    # Prepare the response
    response = {
        "taskId": task_id,  # Assuming task_id is defined somewhere
        "data": data,
    }
    error_code = {"status": StatusCodes.SUCCESS, "reason": "success"}
    response_data = response_template(
        requestId, trace_id, process_duration, response, error_code
    )
    return response_data


def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as file:
        pdf = PdfReader(file)
        text = ""
        for page in range(len(pdf.pages)):
            text += pdf.pages[page].extract_text()
    return {"filename": pdf_path, "text": text}


def extract_text_from_docx(docx_path):
    text = ""
    doc = Document(docx_path)
    for para in doc.paragraphs:
        text += para.text + "\n"
    print(f"Extracted text from {docx_path}: {text}")  # Debugging line
    return {"filename": docx_path, "text": text}


############### CHECK IF ALL INFORMATION IS IN REQUEST ###############
def check_input_request(request):
    reason = ""
    status = ""
    user_id = request.headers.get("X-User-ID", None)
    if user_id is None or not user_id.strip():
        status = StatusCodes.INVALID_REQUEST
        reason = "userToken is invalid"
    request_id = request.headers.get("x-request-id", None)
    request_data = request.get_json()
    print(request_data)
    respose_data = None

    method = request_data["method"]
    print(method)
    if request_id is None or not request_id.strip():
        status = StatusCodes.INVALID_REQUEST
        reason = "requestId is invalid"
    if method is None or not method.strip():
        status = StatusCodes.INVALID_REQUEST
        reason = "method is invalid"
    elif method not in SUPPORTED_METHOD:
        status = StatusCodes.UNSUPPORTED
        reason = f"unsupported method {method}"
    if status != "":
        trace_id = uuid4().hex
        error_code = {"status": status, "reason": reason}
        respose_data = response_template(request_id, trace_id, -1, {}, error_code)

    return respose_data


############### API ENDPOINT TO RECEIVE REQUEST ###############
@app.route("/call", methods=["POST"])
def call_endpoint():
    user_id = request.headers.get("X-User-ID", None)

    request_data = request.get_json()
    task_id = str(uuid.uuid4())
    requestId = str(uuid.uuid4())
    trace_id = str(uuid.uuid4())
    method = request_data["method"]
    print(request)
    if method == "upload_files":
        files = request_data.get("payload").get("files")
        if user_id is not None and files is not None:
            # if user_id is not None and structures is not None and project_name is not None:
            folder_path = f"{user_id}/"
            start_time = time.time()
            request_data = request.get_json()
            project_path = folder_path + "/"

            # Adjusting the pre_signed_urls generation to include folder_path as a prefix
            pre_signed_urls = {
                project_path + file: create_presigned_url(project_path + file)
                for file in files
            }
            data = {"pre_signed_urls": pre_signed_urls}
            end_time = time.time()
            process_duration = end_time - start_time
            response_data = success_response(
                task_id, data, requestId, trace_id, process_duration
            )
            return response_data
        else:
            response = {}
            error_code = {"status": StatusCodes.ERROR, "reason": "User ID not found"}
            response_data = response_template(
                requestId, trace_id, -1, response, error_code
            )
            return response_data

    if method == "process_and_extract":
        # Database insertion
        ret = check_input_request(request)
        if ret is not None:
            return ret

        # Response preparation
        response = {"taskId": task_id}
        error_code = {"status": StatusCodes.PENDING, "reason": "Pending"}
        respose_data = response_template(requestId, trace_id, -1, response, error_code)
        task_status = process_task(task_id,user_id, request_data)
        return task_status
    return True


############### PROCESS THE CALL TASK HERE ###############
def process_task(task_id, user_id, request_data):
    download_url, delete_files_url, processing_duration = process_query(user_id)
    data = {
        "s3_presigned_download": download_url,
        "s3_presigned_delete" : delete_files_url
    }
    # Send the callback
    callback = send_callback(task_id, processing_duration, data)
    return callback


############### SEND CALLBACK TO YOUR APP MARKETPLACE ENDPOINT WITH TASK RESPONSE ###############
def send_callback(task_id, processing_duration, data):

    callback_message = {
        "apiVersion": API_VERSION,
        "service": SERVICE_NAME,
        "datetime": datetime.datetime.now().isoformat(),
        "processDuration": processing_duration,  # Simulated duration
        "taskId": task_id,
        "response": {"dataType": "string", "data": data},
        "StatusCodes": {"status": "TA_000", "reason": "success"},
    }

    headers = {"Content-Type": "application/json"}
    return callback_message
    # response = requests.post(webhook_url, json=callback_message, headers=headers)


############### RUN YOUR SERVER HERE ###############
if __name__ == "__main__":
    app.run()
